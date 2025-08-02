import os
import cv2
import math
import argparse
from pytubefix import YouTube
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
from docx.shared import Inches

def format_time(seconds):
    """Formats seconds into HH:MM:SS string."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02}:{m:02}:{s:02}"

def youtube_to_document(youtube_url: str, output_filename: str = 'youtube_document.docx', screenshot_interval_seconds: int = 60):
    """
    Fetches a YouTube video's details, transcript, and screenshots,
    then compiles them into a Word document.

    Args:
        youtube_url (str): The URL of the YouTube video.
        output_filename (str): The name of the output .docx file.
        screenshot_interval_seconds (int): The interval in seconds between screenshots.
    """
    temp_video_path = "temp_video.mp4"
    temp_screenshot_folder = "temp_screenshots"
    
    try:
        # --- 1. Fetch YouTube Video Details ---
        print("Step 1/5: Fetching video details...")
        yt = YouTube(youtube_url)
        video_id = yt.video_id
        video_title = yt.title
        video_description = yt.description

        # --- 2. Get Transcript ---
        print("Step 2/5: Fetching transcript...")
        try:
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
        except Exception as e:
            print(f"Could not retrieve transcript: {e}")
            transcript_list = [{"text": "Transcript not available for this video.", "start": 0}]

        # --- 3. Download Video for Screenshotting ---
        print("Step 3/5: Downloading video for processing...")
        stream = yt.streams.filter(file_extension="mp4", progressive=True, res="720p").first()
        if not stream:
            stream = yt.streams.filter(file_extension="mp4", progressive=True).first() # Fallback
        stream.download(filename=temp_video_path)
        
        # --- 4. Capture and Save Screenshots ---
        print("Step 4/5: Capturing screenshots...")
        if not os.path.exists(temp_screenshot_folder):
            os.makedirs(temp_screenshot_folder)
            
        screenshot_paths = {}
        cap = cv2.VideoCapture(temp_video_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        video_duration = len(transcript_list) * 2 if not transcript_list[-1]['start'] else transcript_list[-1]['start'] # rough estimate
        
        for i in range(0, int(video_duration), screenshot_interval_seconds):
            frame_id = int(i * fps)
            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_id)
            ret, frame = cap.read()
            if ret:
                timestamp_str = format_time(i)
                screenshot_path = os.path.join(temp_screenshot_folder, f"screenshot_{i}.png")
                cv2.imwrite(screenshot_path, frame)
                screenshot_paths[i] = screenshot_path
        
        cap.release()

        # --- 5. Assemble the Word Document ---
        print("Step 5/5: Assembling the Word document...")
        doc = Document()
        doc.add_heading(video_title, level=1)
        doc.add_paragraph(f"Source URL: {youtube_url}\n")
        
        doc.add_heading("Video Description", level=2)
        doc.add_paragraph(video_description or "No description provided.")
        
        doc.add_heading("Transcript & Screenshots", level=2)

        last_screenshot_time = -1
        for item in transcript_list:
            start_time = item['start']
            
            # Find the closest preceding screenshot to insert
            current_screenshot_time_marker = math.floor(start_time / screenshot_interval_seconds) * screenshot_interval_seconds
            
            if current_screenshot_time_marker > last_screenshot_time:
                if current_screenshot_time_marker in screenshot_paths:
                    doc.add_paragraph(f"\n--- Screenshot at {format_time(current_screenshot_time_marker)} ---\n").bold = True
                    doc.add_picture(screenshot_paths[current_screenshot_time_marker], width=Inches(6.0))
                    last_screenshot_time = current_screenshot_time_marker

            # Add transcript text
            formatted_timestamp = format_time(start_time)
            doc.add_paragraph(f"[{formatted_timestamp}] {item['text']}")
            
        doc.save(output_filename)
        print(f"\nSuccess! Document saved as '{output_filename}'")

    finally:
        # --- Cleanup ---
        print("Cleaning up temporary files...")
        if os.path.exists(temp_video_path):
            os.remove(temp_video_path)
        if os.path.exists(temp_screenshot_folder):
            for f in os.listdir(temp_screenshot_folder):
                os.remove(os.path.join(temp_screenshot_folder, f))
            os.rmdir(temp_screenshot_folder)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert a YouTube video to a Word document with transcript and screenshots.")
    parser.add_argument("url", type=str, help="The full URL of the YouTube video.") # Changed back to positional, removed default for required arg
    parser.add_argument("-o", "--output", type=str, default="开发原生APP-by-AI.docx", help="Name of the output Word file.")
    parser.add_argument("-i", "--interval", type=int, default=30, help="Interval in seconds for taking screenshots.")
    
    args = parser.parse_args()
    
    youtube_to_document(args.url, args.output, args.interval)